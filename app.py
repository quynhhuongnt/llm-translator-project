import streamlit as st
import time
import torch
from transformers import AutoTokenizer, AutoModelForSeq2SeqLM
import easyocr
from PIL import Image
import numpy as np
import docx
import PyPDF2
from io import StringIO, BytesIO

# 1. CẤU HÌNH TRANG
st.set_page_config(page_title="Deep Learning Translator", layout="wide", page_icon="🇬🇧🇻🇳")

# 2. CSS CUSTOM
st.markdown("""
<style>
    .stTextArea textarea { 
        font-size: 16px; 
        height: 300px; 
        font-family: sans-serif;
    }
    
    .stButton button { 
        background-color: #1a73e8; 
        color: white; 
        font-size: 16px; 
        border-radius: 8px; 
        padding: 0.5rem 1rem; 
        border: none; 
        width: 100%; 
        font-weight: bold;
    }
    .stButton button:hover { background-color: #1557b0; color: white; }
    
    .result-box { 
        border: 1px solid #d3d3d3; 
        border-radius: 0.5rem;      
        padding: 1rem;             
        height: 300px;              
        background-color: #f0f2f6;  
        color: #31333F;           
        overflow-y: auto;         
        font-family: sans-serif;   
        font-size: 16px;          
        white-space: pre-wrap;      
        user-select: text;          
        cursor: text;               
    }
    
    header {visibility: hidden;}
    footer {visibility: hidden;}
    
    .lang-header {
        font-weight: bold;
        font-size: 18px;
        margin-bottom: 10px;
        display: block;
        color: #1a73e8;
    }
</style>
""", unsafe_allow_html=True)

# 3. BACKEND - XỬ LÝ MÔ HÌNH VÀ LOGIC

@st.cache_resource
def load_models():
    # Load Model Dịch (EnViT5)
    model_name = "VietAI/envit5-translation"
    tokenizer = AutoTokenizer.from_pretrained(model_name)
    model = AutoModelForSeq2SeqLM.from_pretrained(model_name)
    device = "cuda" if torch.cuda.is_available() else "cpu"
    model = model.to(device)
    
    # Load Model OCR (EasyOCR)
    reader = easyocr.Reader(['en'], gpu=(device == "cuda"))
    
    return tokenizer, model, reader, device

try:
    tokenizer, model, reader, device = load_models()
except Exception as e:
    st.error(f"Lỗi tải model: {e}")

def translate_text(text):
    if not text or text.strip() == "": return ""
    input_text = "en: " + text
    inputs = tokenizer(input_text, return_tensors="pt", padding=True, truncation=True, max_length=512).to(device)
    with torch.no_grad():
        output_ids = model.generate(
            inputs.input_ids,
            max_length=512,
            num_beams=4,
            early_stopping=True
        )
    return tokenizer.decode(output_ids[0], skip_special_tokens=True).replace("vi: ", "")

def split_and_translate(text):
    chunks = text.split('\n')
    translated_chunks = []
    temp_chunk = ""
    for chunk in chunks:
        # Giới hạn token để tránh quá tải mô hình (khoảng 500 ký tự mỗi lần dịch)
        if len(temp_chunk) + len(chunk) < 500:
            temp_chunk += chunk + "\n"
        else:
            translated_chunks.append(translate_text(temp_chunk))
            temp_chunk = chunk + "\n"
    if temp_chunk:
        translated_chunks.append(translate_text(temp_chunk))
    return "\n".join(translated_chunks)

def read_file(uploaded_file):
    text = ""
    if uploaded_file.type == "text/plain":
        stringio = StringIO(uploaded_file.getvalue().decode("utf-8"))
        text = stringio.read()
    elif uploaded_file.type == "application/pdf":
        pdf_reader = PyPDF2.PdfReader(uploaded_file)
        for page in pdf_reader.pages:
            extracted = page.extract_text()
            if extracted: text += extracted + "\n"
    elif "word" in uploaded_file.type:
        doc = docx.Document(uploaded_file)
        for para in doc.paragraphs:
            text += para.text + "\n"
    return text

def create_docx(text):
    """Hàm tạo file Word từ nội dung văn bản"""
    doc = docx.Document()
    for line in text.split('\n'):
        doc.add_paragraph(line)
    
    bio = BytesIO()
    doc.save(bio)
    return bio.getvalue()

# 4. FRONTEND - GIAO DIỆN NGƯỜI DÙNG

st.title(" ỨNG DỤNG DỊCH ANH - VIỆT SỬ DỤNG MÔ HÌNH LLM ")
st.markdown("**Môn:** Kĩ thuật học sâu và ứng dụng")
st.markdown("**Sinh viên thực hiện:** Ngô Thị Quỳnh Hương | **MSV:** 99048")

tab_text, tab_image, tab_doc = st.tabs(["🔤 Văn Bản", "📸 Hình Ảnh", "📂 Tài Liệu"])

# Khởi tạo session state
if 'trans_text' not in st.session_state: st.session_state.trans_text = ""
if 'trans_img' not in st.session_state: st.session_state.trans_img = ""
if 'trans_doc' not in st.session_state: st.session_state.trans_doc = ""

# TAB 1: DỊCH VĂN BẢN TRỰC TIẾP
with tab_text:
    st.write("") 
    c1, c2 = st.columns(2)
    with c1:
        st.markdown('<span class="lang-header">TIẾNG ANH</span>', unsafe_allow_html=True)
        text_input = st.text_area("Input", height=300, placeholder="Nhập văn bản tiếng Anh tại đây...", label_visibility="collapsed")
        
    with c2:
        st.markdown('<span class="lang-header">TIẾNG VIỆT</span>', unsafe_allow_html=True)
        result_content = st.session_state.trans_text if st.session_state.trans_text else ""
        st.markdown(f'<div class="result-box">{result_content}</div>', unsafe_allow_html=True)

    st.write("")
    if st.button("DỊCH VĂN BẢN", key="btn_text"):
        if text_input:
            with st.spinner("AI đang dịch..."):
                st.session_state.trans_text = translate_text(text_input)
                st.rerun()

# TAB 2: DỊCH QUA HÌNH ẢNH (OCR)
with tab_image:
    st.write("")
    c1, c2 = st.columns(2)
    with c1:
        st.markdown('<span class="lang-header">TẢI ẢNH LÊN</span>', unsafe_allow_html=True)
        uploaded_img = st.file_uploader("", type=['png', 'jpg', 'jpeg'], key="upload_img", label_visibility="collapsed")
        if uploaded_img:
            image = Image.open(uploaded_img)
            st.image(image, caption="Ảnh gốc", use_container_width=True)

    with c2:
        st.markdown('<span class="lang-header">KẾT QUẢ DỊCH</span>', unsafe_allow_html=True)
        content_img = st.session_state.trans_img if st.session_state.trans_img else ""
        st.markdown(f'<div class="result-box">{content_img}</div>', unsafe_allow_html=True)
    
    st.write("")
    if st.button("QUÉT & DỊCH ", key="btn_img"):
        if uploaded_img:
            with st.spinner("Đang nhận diện chữ và dịch..."):
                img_np = np.array(image)
                res = reader.readtext(img_np, detail=0)
                extracted_text = " ".join(res)
                translated = translate_text(extracted_text)
                st.session_state.trans_img = f"VĂN BẢN NHẬN DIỆN:\n{extracted_text}\n\nBẢN DỊCH:\n{translated}"
                st.rerun()

# TAB 3: DỊCH TÀI LIỆU VÀ TẢI VỀ .DOCX
with tab_doc:
    st.write("")
    c1, c2 = st.columns(2)
    with c1:
        st.markdown('<span class="lang-header">TẢI FILE (WORD/PDF/TXT)</span>', unsafe_allow_html=True)
        uploaded_doc = st.file_uploader("", type=['docx', 'pdf', 'txt'], key="upload_doc", label_visibility="collapsed")
        if uploaded_doc:
            st.success(f"Đã nhận file: {uploaded_doc.name}")

    with c2:
        st.markdown('<span class="lang-header">NỘI DUNG DỊCH</span>', unsafe_allow_html=True)
        content_doc = st.session_state.trans_doc if st.session_state.trans_doc else "Kết quả sẽ hiện ở đây..."
        display_text = content_doc[:2000] + ("..." if len(content_doc) > 2000 else "")
        st.markdown(f'<div class="result-box">{display_text}</div>', unsafe_allow_html=True)

    st.write("")
    if st.button("DỊCH TOÀN BỘ TÀI LIỆU", key="btn_doc"):
        if uploaded_doc:
            with st.spinner("Đang xử lý file lớn, vui lòng đợi..."):
                raw_text = read_file(uploaded_doc)
                full_translated_text = split_and_translate(raw_text)
                st.session_state.trans_doc = full_translated_text
                st.rerun()
    
    # Nút tải về dạng .docx
    if st.session_state.trans_doc:
        docx_file = create_docx(st.session_state.trans_doc)
        st.download_button(
            label="📄 Tải bản dịch (.docx)",
            data=docx_file,
            file_name="translated_document.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
